#!/usr/bin/env python3
# -*- coding:utf-8 -*-
# ------------------------------------------------------------
# File: docx_draw_host.py
# Created Date: 2020/6/24
# Created Time: 0:14
# Author: Hypdncy
# Author Mail: hypdncy@outlook.com
# Copyright (c) 2020 Hypdncy
# ------------------------------------------------------------
#                       .::::.
#                     .::::::::.
#                    :::::::::::
#                 ..:::::::::::'
#              '::::::::::::'
#                .::::::::::
#           '::::::::::::::..
#                ..::::::::::::.
#              ``::::::::::::::::
#               ::::``:::::::::'        .:::.
#              ::::'   ':::::'       .::::::::.
#            .::::'      ::::     .:::::::'::::.
#           .:::'       :::::  .:::::::::' ':::::.
#          .::'        :::::.:::::::::'      ':::::.
#         .::'         ::::::::::::::'         ``::::.
#     ...:::           ::::::::::::'              ``::.
#    ````':.          ':::::::::'                  ::::..
#                       '.:::::'                    ':'````..
# ------------------------------------------------------------

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from modle.docx_draw import DocxDraw

from cnf.default import table_host_ips
from cnf.data import hostscan_loops
from cnf.const import risk_score
from config import config


class DocxDrawHost(DocxDraw):

    def __init__(self, doc):
        super().__init__(doc)

    def draw(self):
        pass
        self.draw_host_ips()
        self.draw_loopholes()

    def draw_host_ips(self):
        """
        画表
        :return:
        """
        for table in self.doc.tables:
            # 按照列顺序排序
            if len(table.columns) < len(table_host_ips):
                continue
            for col, header in enumerate(table_host_ips):
                if table.cell(0, col).text != header:
                    break
            else:
                num = 0
                for key, value in config["systems"].items():
                    num += 1
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(num)
                    row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row_cells[1].text = key
                    row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row_cells[2].text = value
                    row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                break

    def draw_loopholes(self):
        """
        画漏洞
        :return:
        """
        loops = hostscan_loops.values()
        loops = sorted(loops, key=lambda x: risk_score[x["risk_lev"]], reverse=True)
        for loop in loops:
            self.draw_loophole(loop)

    def draw_loophole(self, loophole):
        """
        画漏洞信息
        :return:
        """

        paragraph0 = self.doc.add_paragraph(
            '【{risk_lev}】{name_cn}'.format(risk_lev=loophole["risk_lev"], name_cn=loophole["name_cn"]))
        paragraph0.style = '安恒信息--标题 2'

        paragraph1_1 = self.doc.add_paragraph('漏洞描述：')
        paragraph1_1.style = '安恒信息--列表（符号二级）'
        paragraph1_2 = self.doc.add_paragraph("{describe}".format(describe=loophole["describe"]))
        paragraph1_2.style = '安恒信息--列表（无符号二级）'

        paragraph2_1 = self.doc.add_paragraph('受影响主机：')
        paragraph2_1.style = '安恒信息--列表（符号二级）'

        table = self.doc.add_table(rows=len(loophole["host_ports"]) + 1, cols=2, style='安恒信息--表格（影响主机）')

        def write_table_rows(row_idx, row_datas):
            row = table.row_cells(row_idx)
            row[0].text, row[1].text = row_datas
            row[0].paragraphs[0].style.font.size = row[1].paragraphs[0].style.font.size = Pt(9)
            row[0].paragraphs[0].alignment = row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        write_table_rows(0, ("主机", "端口"))
        for idx, host_port in enumerate(loophole["host_ports"]):
            write_table_rows(idx + 1, host_port.split('_'))

        paragraph3_1 = self.doc.add_paragraph('加固建议：')
        paragraph3_1.style = '安恒信息--列表（符号二级）'
        paragraph3_2 = self.doc.add_paragraph("{solution}".format(solution=loophole["solution"]))
        paragraph3_2.style = '安恒信息--列表（无符号二级）'
